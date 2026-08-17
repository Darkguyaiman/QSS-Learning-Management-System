from __future__ import annotations

import math
import os
import textwrap
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.pdfbase.pdfmetrics import stringWidth
from reportlab.pdfgen import canvas


ROOT = Path(r"C:\Users\Mohamed Aiman\Desktop\Quick Stop Solution LMS")
REFERENCE = ROOT / "Example" / "System Per Tab Explaination .docx"
WORK = ROOT / "tmp" / "document_generation"
OUT = ROOT / "Generated Task Management System Documentation"
OUT.mkdir(parents=True, exist_ok=True)
WORK.mkdir(parents=True, exist_ok=True)

FONT = "Times New Roman"
BODY_SIZE = 12
NAVY = "111A33"
BLUE = "DCE7FF"
BLUE_STROKE = "6186FF"
MINT = "D8F7E7"
MINT_STROKE = "2DB985"
AMBER = "FFF1C7"
AMBER_STROKE = "E7A52B"
ROSE = "FFE2E2"
ROSE_STROKE = "F06A6A"
VIOLET = "F1E5FF"
VIOLET_STROKE = "9A63E8"
SLATE = "EDF2F7"
SLATE_STROKE = "7C8DA5"


def set_run_font(run, name=FONT, size=BODY_SIZE, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths_inches):
    total_dxa = int(round(sum(widths_inches) * 1440))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_dxa = int(round(widths_inches[idx] * 1440))
            cell.width = Inches(widths_inches[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(header), size=10.5, bold=True)
        set_cell_shading(cell, "DDE5EF")
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(str(value)), size=10.5)
    apply_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return table


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 3), ("Heading 3", 12, 10, 2)):
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st.font.size = Pt(BODY_SIZE)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.0


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* MERGEFORMAT "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=10)


def add_toc(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update the table of contents."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])
    set_run_font(run)


def add_body_paragraph(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        set_run_font(p.add_run(item))


def find_abstract_num_id(doc, paragraph_style="List Number"):
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract.findall(qn("w:lvl")):
            p_style = lvl.find(qn("w:pStyle"))
            if p_style is not None and p_style.get(qn("w:val")) == paragraph_style.replace(" ", ""):
                return abstract.get(qn("w:abstractNumId"))
    return "0"


def new_numbering_id(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), find_abstract_num_id(doc))
    num.append(abstract_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return num_id


def add_steps(doc, items):
    num_id = new_numbering_id(doc)
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num = OxmlElement("w:numId")
        num.set(qn("w:val"), str(num_id))
        num_pr.extend([ilvl, num])
        set_run_font(p.add_run(item))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_run_font(p.add_run(text), size=16 if level == 1 else 13, bold=True)
    return p


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Inches(3.25), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(p.add_run("\tQ-Learning Management System User Guide\t"), size=10, color="52637A")
    add_page_field(p)


def build_docx():
    logo_path = WORK / "qss-logo.png"
    with zipfile.ZipFile(REFERENCE) as zf:
        logo_path.write_bytes(zf.read("word/media/image1.png"))

    doc = Document()
    style_document(doc)
    add_footer(doc.sections[0])

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(230)
    cover_logo = cover.add_run().add_picture(str(logo_path), width=Inches(3.67))
    cover_logo._inline.docPr.set("descr", "Quick Stop Solution logo")
    cover_logo._inline.docPr.set("title", "Quick Stop Solution")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(84)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("Q-Learning Management System"), size=22, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(8)
    set_run_font(subtitle.add_run("Training Task Management System"), size=13, italic=True, color="52637A")
    doc.add_page_break()

    toc_title = doc.add_paragraph()
    toc_title.paragraph_format.space_after = Pt(8)
    set_run_font(toc_title.add_run("Table of Contents"), size=18, bold=True)
    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    add_toc(toc)
    doc.add_page_break()

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    set_run_font(intro.add_run("Tab Explanation and Step-by-Step User Guide"), size=16)
    add_body_paragraph(doc, "This guide explains each main tab and training workspace tab in the Q-Learning Management System, what information appears there, who normally uses it, and the exact steps for common work. It is written for Admin, Trainer, and Trainee users.")
    add_body_paragraph(doc, "Reference style: module purpose, information displayed, main actions, workflow, permissions, and controls. The guide follows the implemented application routes and role checks.")

    add_heading(doc, "System Flow", 1)
    add_body_paragraph(doc, "End-to-end process: Sign in -> Review role dashboard -> Create or open a training -> Assign trainers, healthcare centres, devices, and trainees -> Publish learning content and tests -> Record attendance and practical outcomes -> Complete assessments -> Review and release results -> Issue or download certificates -> Track recertification reminders", "End-to-end process:")

    add_heading(doc, "Sidebar Structure Overview", 1)
    add_body_paragraph(doc, "This navigation map shows the left sidebar used by staff and the simpler navigation available to trainees.")
    tree_lines = [
        "Left Sidebar / Navigation",
        "|-- Dashboard",
        "|   |-- Staff: filters, KPIs, training status, reminders, registrations",
        "|   `-- Trainee: quick actions, learning progress, certificates, recent trainings",
        "|-- Trainings",
        "|   |-- Training list and filters",
        "|   |-- Create or import training, staff only",
        "|   `-- Training workspace tabs",
        "|       |-- Stream / Materials / Tests / Course / People",
        "|       |-- My Results, trainee when released",
        "|       |-- Marks / Attendance / Package, staff",
        "|       |-- Practical Learning Outcomes, Admin for main training",
        "|       `-- Settings, staff",
        "|-- Trainees, Admin/Trainer",
        "|-- Question Bank, Admin/Trainer",
        "|-- Settings, Admin/Trainer",
        "|   |-- Modules / Training Titles / Objectives / Healthcare",
        "|   |-- Areas of Specialization / Device Serial Numbers / Device Models",
        "|   |-- Practical Learning Outcomes",
        "|   `-- Users, Admin only",
        "`-- Profile / Logout"
    ]
    for idx, line in enumerate(tree_lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18 if idx else 0)
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(line), size=10.5, bold=(idx == 0 or line.startswith("|--") or line.startswith("`--")))

    doc.add_page_break()

    add_table(doc,
              ["Sidebar tab", "Sub pages/actions", "Access note"],
              [
                  ["Dashboard", "Role-specific KPIs, filters, quick actions, reminders, progress, certificates.", "All roles."],
                  ["Trainings", "Browse/filter trainings; staff create/manage; open the training workspace.", "All roles; records are filtered by assignment/enrollment."],
                  ["Trainees", "Search, view, create, edit, import, and bulk account actions.", "Admin and Trainer."],
                  ["Question Bank", "Search/filter, create, edit, delete, download template, bulk upload.", "Admin and Trainer."],
                  ["Settings", "Manage LMS reference data and staff accounts.", "Admin and Trainer; Users is Admin only."],
                  ["Profile", "Personal details, profile image, password, and staff certificate upload.", "All roles."],
              ], [1.35, 3.65, 1.5])

    add_table(doc,
              ["Main area", "Purpose", "Typical users"],
              [
                  ["Dashboard", "Daily view of training activity, progress, reminders, and shortcuts.", "All roles"],
                  ["Trainings", "Central list for the full training lifecycle.", "Admin, Trainer, Trainee"],
                  ["Training Workspace", "Content, people, tests, marks, attendance, packages, and settings for one training.", "Role-dependent"],
                  ["Trainees", "Participant account and training-history administration.", "Admin, Trainer"],
                  ["Question Bank", "Reusable assessment questions grouped by learning context.", "Admin, Trainer"],
                  ["Settings", "Reference lists used by registration, training creation, testing, and certification.", "Admin, Trainer"],
                  ["Profile", "Self-service identity, security, and professional information.", "All roles"],
              ], [1.55, 3.5, 1.45])

    doc.add_page_break()
    add_heading(doc, "Role Access", 1)
    add_body_paragraph(doc, "Access is controlled by the role stored in the signed-in session. Training visibility is also limited by assignment, enrollment, status, and lock state.")
    add_table(doc,
              ["Role", "Can view", "Can change"],
              [
                  ["Admin", "All trainings, staff dashboard, trainees, question bank, settings, profiles, reports, and certificates.", "Create/import/manage trainings, grade and release, manage attendance, reference data, and staff users."],
                  ["Trainer", "Assigned in-progress/completed/rescheduled trainings, staff dashboard, trainees, question bank, non-user settings, and profile.", "Create trainings, manage assigned training content, attendance, grading, packages, trainees, questions, and allowed reference data."],
                  ["Trainee", "Own dashboard, enrolled unlocked trainings, permitted materials/tests, own attendance/results, certificates, and profile.", "Update allowed profile fields, upload a profile picture, change password, and complete available tests."],
              ], [1.0, 2.8, 2.7])

    add_heading(doc, "1. Dashboard", 1)
    add_body_paragraph(doc, "Purpose: The Dashboard is the starting point after sign-in. Staff receive an operational reporting view, while trainees receive a personal learning-progress view.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_bullets(doc, [
        "Staff: total, in-progress, completed, canceled, and rescheduled trainings; completion and assessment pass rates.",
        "Staff: dashboard date, healthcare, and module filters, plus downloadable dashboard PDF.",
        "Staff: top clients, healthcare training reminders, upcoming recertifications, trainee status, trainers, and recent registrations.",
        "Trainee: quick links to trainings and profile, learning-progress metrics, certificates, and recent enrolled trainings.",
        "Training cards show current status and the next relevant action for the signed-in user."
    ])
    add_heading(doc, "How To Review Staff Activity", 2)
    add_steps(doc, [
        "Open Dashboard after signing in as Admin or Trainer.",
        "Set the date range and, when needed, filter by healthcare centre or module.",
        "Review training totals, completion rate, assessment pass rate, and the status breakdown.",
        "Check healthcare reminders and upcoming recertifications for follow-up work.",
        "Use Quick Actions to open Trainings, Trainees, or Question Bank.",
        "Click Download PDF when a formal dashboard summary is required."
    ])
    add_heading(doc, "How A Trainee Reviews Progress", 2)
    add_steps(doc, [
        "Open Dashboard after signing in.",
        "Review Learning Progress and the checklist on each recent training.",
        "Click Continue or Review to open an enrolled training.",
        "Open Results when scores have been released.",
        "Use My Certificates to open an available certificate."
    ])

    add_heading(doc, "2. Trainings", 1)
    add_body_paragraph(doc, "Purpose: Trainings is the central list for creating, finding, opening, and tracking learning programmes. Admin sees all records, Trainer sees assigned eligible records, and Trainee sees enrolled unlocked records.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_bullets(doc, [
        "Training title, type, status, trainers, healthcare centres, dates, and enrollment count.",
        "Filters for status, trainer, healthcare, device, training type, and free-text search.",
        "Pending mark-release notifications for staff when follow-up is required.",
        "Create Training for staff and controlled Excel import for authorized administration.",
        "Default header artwork when a training has no custom header image."
    ])
    add_heading(doc, "How To Find And Open A Training", 2)
    add_steps(doc, [
        "Open Trainings from the sidebar or dashboard.",
        "Use search to find a title, description, or trainer.",
        "Apply status, trainer, healthcare, device, or type filters when the list is large.",
        "Open the matching training card.",
        "If access is denied, confirm that the trainer is assigned or the trainee is enrolled and that the training status and lock state permit access."
    ])
    add_heading(doc, "How To Create A Training", 2)
    add_steps(doc, [
        "Sign in as Admin or Trainer and click Create New Training.",
        "Choose the training title, description, type, module, device model, affiliated company, start date/time, and end date/time.",
        "Confirm the required module, device model, start time, and end time.",
        "Select additional trainers; the creator is added automatically.",
        "Select one or more healthcare centres, then choose trainees from those centres.",
        "Select device serial numbers or add approved custom serial numbers.",
        "Click Create Training, then open its workspace to add course content, tests, and operational settings."
    ])
    add_heading(doc, "Training Status And Locking", 2)
    add_bullets(doc, [
        "Admin can see every status; non-admin visibility is limited to in-progress, completed, and rescheduled training records.",
        "Trainees cannot open a locked training and should use released certificate/result paths when available.",
        "Locking makes the workspace read-only and disables content-changing controls.",
        "Use status changes carefully because they affect who can see the training."
    ])

    add_heading(doc, "3. Training Workspace Tabs", 1)
    add_body_paragraph(doc, "Purpose: Opening a training shows a role-specific workspace. The available tabs change according to the signed-in role, training type, enrollment, score release, and lock state.", "Purpose:")
    add_table(doc,
              ["Tab", "Purpose", "Access"],
              [
                  ["Stream", "Training banner, description, overview, materials, and training media.", "All authorized users"],
                  ["Materials", "Trainee-focused material list and access links.", "Trainee"],
                  ["Tests", "Available pre-test, post-test, certificate enrolment, or refresher assessments.", "Trainee"],
                  ["Course", "Sections, materials, media, tests, results, and attendance shortcuts.", "All authorized users; editing for staff"],
                  ["People", "Assigned trainers and enrolled trainees.", "All authorized users"],
                  ["My Results", "Released marks, objective understanding, attendance, and certificate access.", "Trainee after release"],
                  ["Marks", "Scores, attempts, practical grading, result release, and certificate actions.", "Admin/Trainer"],
                  ["Attendance", "Summary, new attendance session, and session history/details.", "Admin/Trainer"],
                  ["Package", "Generate coordinated training documents, PDFs, spreadsheets, and ZIP package.", "Admin/Trainer"],
                  ["Practical Learning Outcomes", "Define main-training practical evaluation criteria.", "Admin, main training"],
                  ["Settings", "Edit training details, assignments, devices, status, and lock controls.", "Admin/Trainer"],
              ], [1.45, 3.65, 1.4])

    add_heading(doc, "How Staff Manage Course Content", 2)
    add_steps(doc, [
        "Open the training and select Course.",
        "Create a section or import sections from another training.",
        "Add a material as a video, document, image, or link and place it in the intended section.",
        "Set material visibility and any access expiry required by the workflow.",
        "Upload training media when photographs or evidence should appear in the course gallery.",
        "Open Tests or the test configuration within the training workflow to confirm each required assessment is available.",
        "Recheck the trainee view before delivery."
    ])
    add_heading(doc, "How Staff Record Attendance", 2)
    add_steps(doc, [
        "Open the training and select Attendance.",
        "Review the attendance summary and any existing sessions.",
        "Open Add New Attendance, enter the date, start time, and end time.",
        "Mark each trainee Present or Absent and add optional notes.",
        "Use the bulk present/absent controls when appropriate, then save.",
        "Open session details or a trainee attendance record to verify totals and attendance rate."
    ])
    add_heading(doc, "How Staff Grade And Release Results", 2)
    add_steps(doc, [
        "Open Marks and review the best attempt for each required test.",
        "For a main training, grade every Practical Learning Outcome criterion.",
        "Open test answers or all attempts when a score needs verification.",
        "Admin releases main-training scores after the required tests and practical evaluation are complete; refresher release can be performed by authorized staff.",
        "Use the exceptional certificate-release action only when the interface permits it and record the required justification.",
        "Confirm that the trainee can now see My Results and the intended certificate action."
    ])
    add_heading(doc, "How A Trainee Completes A Training", 2)
    add_steps(doc, [
        "Open an enrolled training from Dashboard or Trainings.",
        "Read Stream and Course, then open Materials for the learning resources.",
        "Use Tests to start each available assessment and submit answers before leaving the attempt.",
        "Review People and Attendance when participation information is needed.",
        "After staff release scores, open My Results to review tests, practical outcomes, final grades, and certificate availability.",
        "Download or open the certificate only when eligibility and release rules are satisfied."
    ])

    add_heading(doc, "4. Trainees", 1)
    add_body_paragraph(doc, "Purpose: Trainees is the staff administration tab for participant accounts, identity details, healthcare assignment, device information, registration status, and training history.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_bullets(doc, [
        "Name, trainee ID, email, healthcare, IC/passport, handphone, device details, status, and training history.",
        "Search and filters for quickly finding a participant.",
        "Create, view, and edit actions.",
        "Excel template download and bulk import.",
        "Bulk status changes and bulk delete for selected records."
    ])
    add_heading(doc, "How To Add Or Edit A Trainee", 2)
    add_steps(doc, [
        "Open Trainees and click Create Trainee, or find an existing record and click Edit.",
        "Enter or confirm identity, contact, healthcare, specialization, designation, device, and account-status details.",
        "Check that the email and unique identity fields do not conflict with an existing account.",
        "Save the record.",
        "Open the trainee detail view to confirm training and profile information."
    ])
    add_heading(doc, "How To Import Or Update Many Trainees", 2)
    add_steps(doc, [
        "Download the trainee import template.",
        "Complete the spreadsheet using the template headings and accepted values.",
        "Open Import trainees from Excel and choose the completed file.",
        "Review validation messages and correct rejected rows.",
        "Run the import and verify the new or updated records.",
        "For existing records, select rows and use bulk status only after confirming the scope; use bulk delete with extra care."
    ])

    add_heading(doc, "5. Question Bank", 1)
    add_body_paragraph(doc, "Purpose: Question Bank stores reusable assessment questions for training tests. It supports individual editing and spreadsheet-based bulk preparation.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_bullets(doc, [
        "Question text and answer options.",
        "Correct-answer configuration and the related module/objective context.",
        "Search and filters for locating questions.",
        "Edit and delete actions for existing questions.",
        "Template download and bulk upload with validation feedback."
    ])
    add_heading(doc, "How To Create Or Edit A Question", 2)
    add_steps(doc, [
        "Open Question Bank and click Create Question, or click Edit on an existing question.",
        "Select the relevant learning context such as module and objective.",
        "Enter clear question text and the answer options.",
        "Mark the correct answer exactly as required by the test engine.",
        "Save and return to the list.",
        "Use filters to confirm the question appears in the intended group."
    ])
    add_heading(doc, "How To Bulk Upload Questions", 2)
    add_steps(doc, [
        "Download the Question Bank bulk template.",
        "Enter questions using the supplied columns and accepted reference values.",
        "Upload the completed spreadsheet.",
        "Review validation errors before confirming the import.",
        "After import, search for a sample of the new questions and verify their options and correct answers."
    ])

    add_heading(doc, "6. Settings", 1)
    add_body_paragraph(doc, "Purpose: Settings contains the reference lists used by registration, training creation, assessment grouping, device assignment, practical grading, and staff access.", "Purpose:")
    add_table(doc,
              ["Settings tab", "Purpose", "Important actions"],
              [
                  ["Modules", "Group trainings and questions into learning modules.", "Search, add, edit, delete"],
                  ["Training Titles", "Maintain reusable training names and descriptions.", "Search, add, edit, delete"],
                  ["Objectives", "Maintain learning objectives used by questions and scoring.", "Search, add, edit, delete"],
                  ["Healthcare", "Maintain healthcare centres and reminder details; CRM may also synchronize records.", "Search, add, edit, delete with protected identity fields"],
                  ["Areas of Specialization", "Maintain professional specialization values.", "Search, add, edit, delete"],
                  ["Device Serial Numbers", "Maintain serial numbers linked to device models.", "Search, add, edit, delete"],
                  ["Device Models", "Maintain available medical-device models.", "Search, add, edit, delete"],
                  ["Practical Learning Outcomes", "Maintain reusable criteria for main-training assessment.", "Search, add, edit, delete"],
                  ["Users", "Manage Admin and Trainer accounts and staff certificates.", "Admin only"],
              ], [1.55, 3.25, 1.7])
    add_heading(doc, "How To Maintain Reference Data", 2)
    add_steps(doc, [
        "Open Settings and choose the required reference list.",
        "Search first to avoid creating a duplicate value.",
        "Click Add, complete the required fields, and save.",
        "Use Edit when a label or relationship changes.",
        "Delete only after confirming that the value is not required by existing training or assessment records.",
        "For healthcare records, remember that CRM synchronization and protected identity-field rules may apply."
    ])
    add_heading(doc, "How Admin Manages Staff Users", 2)
    add_steps(doc, [
        "Sign in as Admin and open Settings > Users.",
        "Search for an existing account before creating a new one.",
        "Add or edit the staff member's identity, email, role, professional details, password, and certificate information as required.",
        "Save the account and test the intended role access.",
        "Delete a staff account only after reviewing training assignments and audit implications."
    ])

    add_heading(doc, "7. Profile", 1)
    add_body_paragraph(doc, "Purpose: Profile lets each signed-in user maintain allowed personal information, profile picture, and password. Staff can also store a professional certificate.", "Purpose:")
    add_heading(doc, "Information Displayed", 2)
    add_bullets(doc, [
        "All roles: name, email, role or trainee ID, profile picture, and password controls.",
        "Staff: position, phone number, areas of specialization, and certificate upload/view.",
        "Trainee: healthcare, designation, specialization, device serial number, training dates, recertification date, completed-training count, and account status.",
        "Some trainee identity and assignment fields are read-only in self-service Profile."
    ])
    add_heading(doc, "How To Update Profile And Password", 2)
    add_steps(doc, [
        "Open Profile from the user area.",
        "Update the editable personal or professional fields and click Update Profile.",
        "Use Upload Picture to replace the profile image when needed.",
        "Staff may upload a certificate as a PDF or image within the displayed file-size rule.",
        "For a password change, enter the current password, new password, and confirmation.",
        "Click Change Password and sign in again if the session is ended."
    ])

    add_heading(doc, "8. Tests, Results, Attendance, And Certificates", 1)
    add_body_paragraph(doc, "Purpose: These linked functions record assessment attempts, practical outcomes, attendance, final grades, score release, and certificate eligibility for each enrollment.", "Purpose:")
    add_heading(doc, "Assessment And Result Logic", 2)
    add_bullets(doc, [
        "Main training can include pre-test, post-test, Practical Learning Outcomes, and certificate enrolment assessment.",
        "Refresher training uses its configured assessment path, including certificate enrolment when available.",
        "The system stores completed attempts and uses the best completed score where the result workflow requires it.",
        "Trainees see released results only for their own enrollment.",
        "Certificate availability depends on completed assessments, required thresholds, practical outcomes for main training, lockout rules, and any authorized release override."
    ])
    add_heading(doc, "How To Review One Enrollment", 2)
    add_steps(doc, [
        "Open the training and select Marks, or open the trainee from a result/attendance shortcut.",
        "Review test attempts, objective understanding, Practical Learning Outcomes, and attendance totals.",
        "Confirm final-grade calculations after all required inputs are present.",
        "Release results only when the record is complete and verified.",
        "Open the certificate preview and confirm participant name, training details, and issue information before distribution."
    ])

    add_heading(doc, "System Controls And Good Practice", 1)
    add_heading(doc, "Required Fields And Validation", 2)
    add_bullets(doc, [
        "Training creation requires valid dates, a module, a device model, participating healthcare centre(s), and at least one trainee.",
        "Email and unique identity fields must remain unique across the applicable user table.",
        "Bulk files must keep the provided headings and accepted values.",
        "Attendance sessions require date, start time, end time, and a status for each included trainee.",
        "Result release should wait until all required assessments and practical criteria are complete.",
        "Locked trainings should be treated as read-only."
    ])
    add_heading(doc, "Recommended Operating Routine", 2)
    add_steps(doc, [
        "Start on Dashboard and review reminders, recertifications, pending releases, and recent registrations.",
        "Open Trainings and confirm dates, status, assigned trainer, healthcare centres, devices, and trainees.",
        "Prepare sections, materials, tests, and Practical Learning Outcomes before delivery.",
        "Record attendance for every session and verify exceptions.",
        "Review attempts and practical scores, calculate grades, then release results.",
        "Generate the training package and confirm certificate availability.",
        "Lock or complete the training only after the operational record is finished."
    ])
    add_heading(doc, "Complete Information Flow", 2)
    add_body_paragraph(doc, "LMS flow: Registration or staff-created trainee -> Healthcare, specialization, device, and profile data -> Training creation and enrollment -> Sections, materials, tests, and training media -> Attendance and assessment attempts -> Practical grading and final grades -> Score release -> Certificate and recertification tracking -> Dashboard reporting and reminders", "LMS flow:")

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output = OUT / "System Per Tab Explanation - Q-Learning Management System.docx"
    doc.save(output)
    return output


def hex_rgb(value):
    value = value.lstrip("#")
    return tuple(int(value[i:i+2], 16) / 255 for i in (0, 2, 4))


def wrap_lines(text, max_chars):
    lines = []
    for raw in str(text).split("\n"):
        lines.extend(textwrap.wrap(raw, width=max_chars, break_long_words=False) or [""])
    return lines


class DiagramWriter:
    def __init__(self, width, height, pdf_path, svg_path):
        self.width = width
        self.height = height
        self.pdf = canvas.Canvas(str(pdf_path), pagesize=(width, height))
        self.svg = [f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">',
                    '<defs><marker id="arrow" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto" markerUnits="strokeWidth"><path d="M0,0 L0,6 L9,3 z" fill="#8EA6D9"/></marker></defs>',
                    '<rect width="100%" height="100%" fill="#FFFFFF"/>']

    def rect(self, x, y, w, h, fill, stroke, radius=0, line=1.2):
        self.pdf.setFillColorRGB(*hex_rgb(fill))
        self.pdf.setStrokeColorRGB(*hex_rgb(stroke))
        self.pdf.setLineWidth(line)
        if radius:
            self.pdf.roundRect(x, y, w, h, radius, fill=1, stroke=1)
            self.svg.append(f'<rect x="{x}" y="{self.height-y-h}" width="{w}" height="{h}" rx="{radius}" fill="#{fill}" stroke="#{stroke}" stroke-width="{line}"/>')
        else:
            self.pdf.rect(x, y, w, h, fill=1, stroke=1)
            self.svg.append(f'<rect x="{x}" y="{self.height-y-h}" width="{w}" height="{h}" fill="#{fill}" stroke="#{stroke}" stroke-width="{line}"/>')

    def text(self, x, y, text, size=12, color="334155", bold=False, align="center", max_chars=28, leading=None):
        lines = wrap_lines(text, max_chars)
        leading = leading or size * 1.25
        font = "Helvetica-Bold" if bold else "Helvetica"
        self.pdf.setFont(font, size)
        self.pdf.setFillColorRGB(*hex_rgb(color))
        for i, line in enumerate(lines):
            yy = y - i * leading
            if align == "center":
                self.pdf.drawCentredString(x, yy, line)
                anchor = "middle"
            elif align == "right":
                self.pdf.drawRightString(x, yy, line)
                anchor = "end"
            else:
                self.pdf.drawString(x, yy, line)
                anchor = "start"
            weight = "700" if bold else "400"
            self.svg.append(f'<text x="{x}" y="{self.height-yy}" text-anchor="{anchor}" font-family="Arial, Helvetica, sans-serif" font-size="{size}" font-weight="{weight}" fill="#{color}">{escape(line)}</text>')

    def line(self, x1, y1, x2, y2, color="8EA6D9", width=1.2, arrow=False, dashed=False):
        self.pdf.setStrokeColorRGB(*hex_rgb(color))
        self.pdf.setLineWidth(width)
        if dashed:
            self.pdf.setDash(4, 3)
        self.pdf.line(x1, y1, x2, y2)
        self.pdf.setDash()
        dash = ' stroke-dasharray="4 3"' if dashed else ""
        marker = ' marker-end="url(#arrow)"' if arrow else ""
        self.svg.append(f'<line x1="{x1}" y1="{self.height-y1}" x2="{x2}" y2="{self.height-y2}" stroke="#{color}" stroke-width="{width}"{dash}{marker}/>')
        if arrow:
            angle = math.atan2(y2-y1, x2-x1)
            length = 8
            spread = 0.45
            p1 = (x2-length*math.cos(angle-spread), y2-length*math.sin(angle-spread))
            p2 = (x2-length*math.cos(angle+spread), y2-length*math.sin(angle+spread))
            self.pdf.setFillColorRGB(*hex_rgb(color))
            path = self.pdf.beginPath()
            path.moveTo(x2, y2); path.lineTo(*p1); path.lineTo(*p2); path.close()
            self.pdf.drawPath(path, fill=1, stroke=0)

    def finish(self):
        self.pdf.showPage()
        self.pdf.save()
        self.svg.append('</svg>')
        Path(self.svg_path if hasattr(self, 'svg_path') else '').write_text('', encoding='utf-8') if False else None


def save_diagram(writer, svg_path):
    writer.pdf.showPage()
    writer.pdf.save()
    writer.svg.append("</svg>")
    svg_path.write_text("\n".join(writer.svg), encoding="utf-8")


def role_diagram(role, available, unavailable):
    panels = available + ([{"title": f"Not available for {role}", "nodes": unavailable, "fill": ROSE, "stroke": ROSE_STROKE}] if unavailable else [])
    node_w = 180
    node_h = 58
    panel_gap = 28
    panel_padding = 24
    widths = []
    for p in panels:
        cols = min(3, max(1, math.ceil(len(p["nodes"]) / 2)))
        widths.append(max(300, cols * node_w + (cols-1)*24 + panel_padding*2))
    width = sum(widths) + panel_gap * (len(panels)-1) + 100
    height = 760
    pdf_path = OUT / f"{role} Role User Flow Diagram.pdf"
    svg_path = OUT / f"{role} Role User Flow Diagram.svg"
    d = DiagramWriter(width, height, pdf_path, svg_path)

    title_w, title_h = 260, 82
    d.rect(width/2-title_w/2, height-110, title_w, title_h, NAVY, NAVY)
    d.text(width/2, height-56, f"{role} Role -", size=18, color="FFFFFF", bold=True, max_chars=30)
    d.text(width/2, height-82, "How to Use Each Tab", size=16, color="FFFFFF", bold=True, max_chars=30)
    pill_w = 170
    d.rect(width/2-pill_w/2, height-165, pill_w, 34, SLATE_STROKE if role == "Trainee" else "147D75", SLATE_STROKE if role == "Trainee" else "147D75", radius=17)
    d.text(width/2, height-151, f"Sign in as {role}", size=11, color="FFFFFF", bold=True, max_chars=30)

    x = 50
    sign_y = height-165
    for idx, panel in enumerate(panels):
        w = widths[idx]
        fill = panel.get("fill", [BLUE, MINT, AMBER, VIOLET, SLATE][idx % 5])
        stroke = panel.get("stroke", [BLUE_STROKE, MINT_STROKE, AMBER_STROKE, VIOLET_STROKE, SLATE_STROKE][idx % 5])
        panel_y = 120
        panel_h = 430
        d.rect(x, panel_y, w, panel_h, fill, stroke)
        d.text(x+w/2, panel_y+panel_h-24, panel["title"], size=12, color=stroke, bold=True, max_chars=45)
        d.line(width/2, sign_y, x+w/2, panel_y+panel_h, color="C8D4F3", width=1.0, arrow=True)

        nodes = panel["nodes"]
        cols = min(3, max(1, math.ceil(len(nodes)/2)))
        rows = math.ceil(len(nodes)/cols)
        start_x = x + (w - (cols*node_w + (cols-1)*24))/2
        start_y = panel_y + panel_h - 100
        positions = []
        for n_idx, node in enumerate(nodes):
            row = n_idx // cols
            col = n_idx % cols
            nx = start_x + col*(node_w+24)
            ny = start_y - row*(node_h+42)
            nfill = "FFFFFF" if not node.startswith("NO:") else ROSE
            nstroke = stroke if not node.startswith("NO:") else ROSE_STROKE
            label = node[3:] if node.startswith("NO:") else node
            d.rect(nx, ny-node_h, node_w, node_h, nfill, nstroke)
            d.text(nx+node_w/2, ny-22, label, size=10, color="334155" if not node.startswith("NO:") else "9B3C3C", max_chars=26, leading=12)
            positions.append((nx+node_w/2, ny-node_h/2))
        for a, b in zip(positions, positions[1:]):
            d.line(a[0], a[1]-node_h/2+2, b[0], b[1]+node_h/2-2, color=stroke, width=0.8, arrow=True)
        x += w + panel_gap
    save_diagram(d, svg_path)
    return pdf_path, svg_path


def overview_diagram():
    width, height = 2500, 1380
    pdf_path = OUT / "Full System Overview Diagram Q-Learning Management System.pdf"
    svg_path = OUT / "Full System Overview Diagram Q-Learning Management System.svg"
    d = DiagramWriter(width, height, pdf_path, svg_path)
    d.text(width/2, height-70, "Full System Overview Diagram", size=38, color="111111", max_chars=60)
    d.text(width/2, height-112, "Q-Learning Management System", size=20, color="52637A", max_chars=60)

    nodes = {
        "signin": (80, 1160, 170, 58, "Sign in / Register", "147D75", "147D75"),
        "dash": (80, 1010, 240, 92, "Role Dashboard\nKPIs - progress - reminders", "0D6E9E", "0D6E9E"),
        "trainings": (420, 1010, 260, 92, "Trainings\nCreate, filter, assign, enroll", BLUE, BLUE_STROKE),
        "workspace": (790, 1010, 300, 92, "Training Workspace\nContent - people - tests", MINT, MINT_STROKE),
        "delivery": (1210, 1010, 300, 92, "Delivery Records\nAttendance - attempts - practical", AMBER, AMBER_STROKE),
        "results": (1630, 1010, 300, 92, "Results & Certificates\nGrade - release - issue", VIOLET, VIOLET_STROKE),
        "report": (2050, 1010, 300, 92, "Reporting & Reminders\nDashboard - package - recertification", SLATE, SLATE_STROKE),
        "trainees": (420, 710, 260, 92, "Trainees\nAccounts - identity - history", BLUE, BLUE_STROKE),
        "questions": (790, 710, 300, 92, "Question Bank\nQuestions - options - objectives", MINT, MINT_STROKE),
        "settings": (1210, 710, 300, 92, "Settings\nModules - titles - healthcare - devices", AMBER, AMBER_STROKE),
        "profile": (1630, 710, 300, 92, "Profile\nIdentity - security - certificates", VIOLET, VIOLET_STROKE),
        "crm": (2050, 710, 300, 92, "CRM / Imports\nHealthcare sync - Excel templates", SLATE, SLATE_STROKE),
    }
    for key, (x,y,w,h,label,fill,stroke) in nodes.items():
        text_color = "FFFFFF" if key in ("signin", "dash") else "334155"
        d.rect(x,y,w,h,fill,stroke,radius=18 if key=="signin" else 0,line=1.5)
        d.text(x+w/2,y+h/2+11,label,size=13 if key!="dash" else 14,color=text_color,bold=key in ("signin","dash"),max_chars=32,leading=17)

    def arrow(a,b):
        ax,ay,aw,ah,*_ = nodes[a]; bx,by,bw,bh,*_ = nodes[b]
        d.line(ax+aw, ay+ah/2, bx, by+bh/2, color="1F2937", width=1.8, arrow=True)
    d.line(165,1160,165,1102,color="1F2937",width=1.8,arrow=True)
    for a,b in (("dash","trainings"),("trainings","workspace"),("workspace","delivery"),("delivery","results"),("results","report")):
        arrow(a,b)
    for key in ("trainees","questions","settings","profile","crm"):
        x,y,w,h,*_ = nodes[key]
        d.line(200,1010,x+w/2,y+h,color="1F2937",width=1.3,arrow=True)
    d.line(550,802,550,1010,color=BLUE_STROKE,width=1.2,arrow=True)
    d.line(940,802,940,1010,color=MINT_STROKE,width=1.2,arrow=True)
    d.line(1360,802,1360,1010,color=AMBER_STROKE,width=1.2,arrow=True)

    d.text(210, 500, "Training lifecycle", size=18, color="111111", bold=True, align="left", max_chars=40)
    lifecycle = [
        (80,360,"Draft / Create",BLUE,BLUE_STROKE),
        (360,360,"In Progress",MINT,MINT_STROKE),
        (640,360,"Attendance & Tests",AMBER,AMBER_STROKE),
        (920,360,"Grade & Release",VIOLET,VIOLET_STROKE),
        (1200,360,"Completed / Locked",SLATE,SLATE_STROKE),
        (1480,360,"Certificate",MINT,MINT_STROKE),
        (1760,360,"Recertification",AMBER,AMBER_STROKE),
    ]
    for i,(x,y,label,fill,stroke) in enumerate(lifecycle):
        d.rect(x,y,220,70,fill,stroke)
        d.text(x+110,y+31,label,size=12,color="334155",bold=True,max_chars=24)
        if i < len(lifecycle)-1:
            d.line(x+220,y+35,lifecycle[i+1][0],y+35,color="1F2937",width=1.4,arrow=True)

    d.rect(2040,250,380,340,"F7F9FC",SLATE_STROKE)
    d.text(2230,555,"Who can do what",size=16,color="52637A",bold=True,max_chars=35)
    legend = [
        ("Trainee - learn, test, view own results",SLATE,SLATE_STROKE),
        ("Trainer - assigned training operations",MINT,MINT_STROKE),
        ("Admin - full oversight and user access",ROSE,ROSE_STROKE),
    ]
    ly=480
    for label,fill,stroke in legend:
        d.rect(2080,ly,300,66,fill,stroke)
        d.text(2230,ly+28,label,size=11,color="334155",max_chars=38)
        ly-=92
    save_diagram(d, svg_path)
    return pdf_path, svg_path


def build_diagrams():
    outputs = []
    outputs.extend(role_diagram("Admin", [
        {"title":"Dashboard","nodes":["Review KPIs and filters","Download dashboard PDF","Check reminders and registrations"]},
        {"title":"Trainings","nodes":["Browse all trainings","Create or authorized Excel import","Manage content, tests, status and lock","Review mark-release notifications"]},
        {"title":"Training Workspace","nodes":["Stream / Course / People","Marks and Attendance","Package generation","Practical Learning Outcomes","Settings and assignments"]},
        {"title":"Trainees & Questions","nodes":["Create, edit and import trainees","Bulk status or delete","Create, edit and import questions"]},
        {"title":"Settings & Profile","nodes":["Manage reference lists","Manage Admin/Trainer users","Update profile, certificate and password"]},
    ], []))
    outputs.extend(role_diagram("Trainer", [
        {"title":"Dashboard","nodes":["Review KPIs and filters","Download dashboard PDF","Check reminders and registrations"]},
        {"title":"Trainings","nodes":["Browse assigned eligible trainings","Create a training","Manage content, tests and assignments","Review mark-release notifications"]},
        {"title":"Training Workspace","nodes":["Stream / Course / People","Marks and Attendance","Package generation","Training Settings"]},
        {"title":"Trainees & Questions","nodes":["Create, edit and import trainees","Bulk trainee actions","Create, edit and import questions"]},
        {"title":"Settings & Profile","nodes":["Manage allowed reference lists","Update profile and certificate","Change password and sign out"]},
    ], ["NO:View unassigned or draft-only trainings","NO:Manage staff user accounts","NO:Admin-only main-training release controls"]))
    outputs.extend(role_diagram("Trainee", [
        {"title":"Dashboard","nodes":["Review learning progress","Open recent trainings","View available certificates"]},
        {"title":"Trainings","nodes":["Browse enrolled unlocked trainings","Open Stream and Course","View Materials","Take available Tests"]},
        {"title":"Training Workspace","nodes":["View People","Review own attendance","Open My Results after release","Download eligible certificate"]},
        {"title":"Profile","nodes":["Update allowed details","Upload profile picture","Change password and sign out"]},
    ], ["NO:Create or edit trainings","NO:Manage trainees or questions","NO:Mark attendance or grade","NO:Open Settings or staff packages"]))
    outputs.extend(overview_diagram())
    return outputs


if __name__ == "__main__":
    docx = build_docx()
    diagrams = build_diagrams()
    print(docx)
    for item in diagrams:
        print(item)
