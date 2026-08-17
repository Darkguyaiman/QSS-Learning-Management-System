from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from pathlib import Path

from docx import Document


def rgb(value):
    return str(value) if value is not None else None


path = Path(sys.argv[1]).resolve()
doc = Document(path)

result = {
    "path": str(path),
    "sha256": hashlib.sha256(path.read_bytes()).hexdigest(),
    "sections": [],
    "paragraphs": [],
    "tables": [],
    "headers": [],
    "footers": [],
    "inline_shapes": len(doc.inline_shapes),
}

for section in doc.sections:
    result["sections"].append({
        "page_width": section.page_width,
        "page_height": section.page_height,
        "top_margin": section.top_margin,
        "bottom_margin": section.bottom_margin,
        "left_margin": section.left_margin,
        "right_margin": section.right_margin,
        "header_distance": section.header_distance,
        "footer_distance": section.footer_distance,
        "different_first_page_header_footer": section.different_first_page_header_footer,
    })
    result["headers"].append([p.text for p in section.header.paragraphs])
    result["footers"].append([p.text for p in section.footer.paragraphs])

for idx, p in enumerate(doc.paragraphs):
    runs = []
    for r in p.runs:
        runs.append({
            "text": r.text,
            "bold": r.bold,
            "italic": r.italic,
            "underline": bool(r.underline) if r.underline is not None else None,
            "font": r.font.name,
            "size_pt": r.font.size.pt if r.font.size else None,
            "color": rgb(r.font.color.rgb),
        })
    result["paragraphs"].append({
        "index": idx,
        "text": p.text,
        "style": p.style.name if p.style else None,
        "alignment": p.alignment,
        "left_indent": p.paragraph_format.left_indent,
        "right_indent": p.paragraph_format.right_indent,
        "first_line_indent": p.paragraph_format.first_line_indent,
        "space_before": p.paragraph_format.space_before,
        "space_after": p.paragraph_format.space_after,
        "line_spacing": p.paragraph_format.line_spacing,
        "page_break_before": p.paragraph_format.page_break_before,
        "keep_with_next": p.paragraph_format.keep_with_next,
        "runs": runs,
    })

for t_idx, table in enumerate(doc.tables):
    result["tables"].append({
        "index": t_idx,
        "style": table.style.name if table.style else None,
        "rows": [[cell.text for cell in row.cells] for row in table.rows],
    })

with zipfile.ZipFile(path) as zf:
    result["package_parts"] = [
        {"name": i.filename, "size": i.file_size, "sha256": hashlib.sha256(zf.read(i.filename)).hexdigest()}
        for i in zf.infolist()
    ]

print(json.dumps(result, indent=2, default=str))
