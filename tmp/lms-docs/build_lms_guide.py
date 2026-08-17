from __future__ import annotations

import hashlib
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
REFERENCE = ROOT / "Example" / "System Per Tab Explaination .docx"
OUTPUT_DIR = ROOT / "LMS Documentation"
OUTPUT = OUTPUT_DIR / "Q-Learning Management System Per Tab Explanation.docx"
EXPECTED_SHA = "88270f62762c4a7bffaa8ae44cd37c220c4b66b603152c43591d531f7b7106d9"


def set_font(run, size=11, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_inches):
    widths_dxa = [round(w * 1440) for w in widths_inches]
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def paragraph(doc, text="", size=11, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              before=0, after=6, keep=False, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_with_next = keep
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic)
    return p


def label_paragraph(doc, label, text):
    p = paragraph(doc, after=6)
    r = p.add_run(label)
    set_font(r, 11, True)
    r = p.add_run(text)
    set_font(r, 11)
    return p


def main_heading(doc, text):
    return paragraph(doc, text, size=14, bold=True, before=8, after=5, keep=True)


def subheading(doc, text):
    return paragraph(doc, text, size=11, bold=True, before=5, after=3, keep=True)


def bullets(doc, items):
    for item in items:
        p = paragraph(doc, item, size=10.5, after=2, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)


def steps(doc, items):
    for item in items:
        p = paragraph(doc, item, size=10.5, after=2, style="List Number")
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.18)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_font(run, 9, True)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    set_font(run, 9)
    set_table_geometry(table, widths)
    paragraph(doc, "", after=2)
    return table


def nav_line(doc, text, level=0, bold=False):
    p = paragraph(doc, text, size=8.5, bold=bold, after=1, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.left_indent = Inches(level * 0.16)


def verify_reference():
    digest = hashlib.sha256(REFERENCE.read_bytes()).hexdigest()
    if digest != EXPECTED_SHA:
        raise RuntimeError(f"Reference DOCX changed: {digest}")


verify_reference()
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
shutil.copy2(REFERENCE, OUTPUT)
doc = Document(OUTPUT)
clear_body(doc)

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

for style_name in ("Normal", "List Bullet", "List Number"):
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    style.font.size = Pt(11)

with zipfile.ZipFile(REFERENCE) as source_zip:
    logo_bytes = source_zip.read("word/media/image1.png")
logo_path = ROOT / "tmp" / "lms-docs" / "qss-logo.png"
logo_path.write_bytes(logo_bytes)

paragraph(doc, "", after=1)
cover_logo = doc.add_paragraph()
cover_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_logo.paragraph_format.space_after = Pt(18)
logo_shape = cover_logo.add_run().add_picture(str(logo_path), width=Inches(4.7))
logo_shape._inline.docPr.set("descr", "Quick Stop Solution logo")
logo_shape._inline.docPr.set("title", "Quick Stop Solution")
paragraph(doc, "Q-Learning Management System", size=24, bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
paragraph(doc, "Tab Explanation and Step-by-Step User Guide", size=14,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
paragraph(doc, "Quick Stop Solution", size=11, italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
doc.add_page_break()

paragraph(doc, "This guide explains each LMS tab, the information shown, who normally uses it, and the steps for completing common work. It is written for Admin, Trainer, and Trainee users.")
paragraph(doc, "Reference style: module purpose, information displayed, main actions, workflow, and controls. The guide covers the training lifecycle from registration and enrollment through materials, tests, attendance, grading, package generation, results, and certification.")

subheading(doc, "System Flow")
label_paragraph(doc, "End-to-end process: ", "Trainee registers or staff creates a trainee -> Staff creates and configures a training -> Healthcare, trainers, trainees, and devices are assigned -> Sections, materials, questions, tests, and practical outcomes are prepared -> Training becomes active -> Attendance and test attempts are recorded -> Staff evaluates marks and practical outcomes -> Results are released -> Certificate and training package become available.")

subheading(doc, "Sidebar Structure Overview")
paragraph(doc, "This is the practical navigation map for the LMS sidebar. Items are role-sensitive: Admin and Trainer users see staff-management pages, while Trainees see a simpler learning navigation.")
nav_line(doc, "Left Sidebar", bold=True)
nav_line(doc, "|-- Dashboard", bold=True)
nav_line(doc, "|   |-- Staff operational dashboard and PDF export", 1)
nav_line(doc, "|   `-- Trainee progress, training, results, and certificate overview", 1)
nav_line(doc, "|-- Trainings", bold=True)
nav_line(doc, "|   |-- Search/open training cards", 1)
nav_line(doc, "|   |-- Create training / bulk import, Admin and Trainer", 1)
nav_line(doc, "|   `-- Open the role-sensitive training workspace", 1)
nav_line(doc, "|       |-- Stream", 2, True)
nav_line(doc, "|       |-- Materials and Tests, Trainee", 2)
nav_line(doc, "|       |-- Course and People", 2)
nav_line(doc, "|       |-- My Results, Trainee after release", 2)
nav_line(doc, "|       |-- Marks, Attendance, Package, Admin and Trainer", 2)
nav_line(doc, "|       |-- Practical Learning Outcomes, Admin on main training", 2)
nav_line(doc, "|       `-- Settings, Admin and Trainer", 2)
nav_line(doc, "|-- Trainees, Admin and Trainer", bold=True)
nav_line(doc, "|   |-- Search/filter, create/edit, bulk import", 1)
nav_line(doc, "|   `-- Bulk status update and bulk delete", 1)
nav_line(doc, "|-- Question Bank, Admin and Trainer", bold=True)
nav_line(doc, "|   |-- Search/filter questions", 1)
nav_line(doc, "|   |-- Create/edit/delete", 1)
nav_line(doc, "|   `-- Download template / bulk upload", 1)
nav_line(doc, "|-- Settings, Admin and Trainer", bold=True)
for item in ("Modules", "Training Titles", "Objectives", "Healthcare", "Areas of Specialization", "Device Serial Numbers", "Device Models", "Practical Learning Outcomes"):
    nav_line(doc, f"|   |-- {item}", 1)
nav_line(doc, "|   `-- Users, Admin only", 1)
nav_line(doc, "`-- Profile / Logout", bold=True)

subheading(doc, "Navigation Summary")
add_table(doc, ["Sidebar tab", "Subpages/actions", "Access note"], [
    ("Dashboard", "Role-specific summaries, progress or operational KPIs, date filters, staff PDF export.", "All roles."),
    ("Trainings", "Training cards, create/import, training workspace, lifecycle and learning activities.", "All roles; management for Admin/Trainer."),
    ("Trainees", "Search, filters, create/edit, spreadsheet import, bulk status and deletion.", "Admin and Trainer."),
    ("Question Bank", "Questions, answers, test type, module/objective mapping, bulk upload.", "Admin and Trainer."),
    ("Settings", "Modules, titles, objectives, healthcare, specialties, devices, models, practical outcomes, users.", "Admin and Trainer; Users is Admin only."),
    ("Profile", "Personal information, photo, certificate file where applicable, password change.", "All roles."),
], [1.45, 3.55, 1.5])

subheading(doc, "Role Access")
paragraph(doc, "Access is controlled by the role stored in the signed-in session. The training workspace also hides or shows tabs according to the role, training type, enrollment, and whether results have been released.")
add_table(doc, ["Role", "Can view", "Can change"], [
    ("Admin", "All staff pages, all training-management tabs, reports, reference settings, and user accounts.", "Create and configure trainings; manage trainees and questions; attendance, marks, score release, certificates, packages, practical outcomes, settings, and users."),
    ("Trainer", "Staff dashboard, trainings, trainees, question bank, most training-management tabs, and reference settings.", "Create/configure assigned trainings; manage content, enrollment, attendance, marks, packages, trainees, questions, and non-user settings within route permissions."),
    ("Trainee", "Own dashboard, enrolled trainings, learning materials, tests, people, attendance, released results, and certificates.", "Self-register; update own profile; take tests; access assigned materials; download results/certificates when released."),
], [1.1, 2.65, 2.75])

main_heading(doc, "1. Dashboard")
label_paragraph(doc, "Purpose: ", "The Dashboard is the first role-specific overview after sign-in. Staff use it to monitor training activity and trainee status; trainees use it to understand their own learning progress and available results.")
subheading(doc, "Information Displayed")
bullets(doc, [
    "Admin/Trainer: training totals and statuses, trainee status totals, trainer workload, registrations, top healthcare facilities or client activity, recertification information, and recent certificate activity.",
    "Staff reporting controls for date range and healthcare filtering where available.",
    "A staff PDF export of the dashboard report.",
    "Trainee: profile welcome information, enrolled training cards, training progress, scores or results availability, and issued certificates.",
])
subheading(doc, "How To Review The Staff Dashboard")
steps(doc, [
    "Open Dashboard after signing in as Admin or Trainer.",
    "Set the date range or healthcare filter when the review should focus on a specific period or facility.",
    "Review training status, trainee status, recent registrations, trainer assignments, and recertification indicators.",
    "Open the relevant Training or Trainees page to act on a record.",
    "Use the PDF action when a fixed management summary is required.",
])
subheading(doc, "How A Trainee Uses The Dashboard")
steps(doc, [
    "Review the welcome/profile block and current enrolled trainings.",
    "Open a training card to continue materials, coursework, or tests.",
    "Check whether results or certificates have been released.",
    "Open My Results or the certificate download from the relevant training when available.",
])

main_heading(doc, "2. Trainings")
label_paragraph(doc, "Purpose: ", "Trainings is the central list for learning programmes. Staff create and manage training records; trainees open the trainings in which they are enrolled.")
subheading(doc, "Information Displayed")
bullets(doc, [
    "Training title, type, module, schedule, status, assigned trainer, healthcare facilities, and enrollment context.",
    "Training cards that open the detailed training workspace.",
    "Create New Training and Excel template/import actions for Admin and Trainer users.",
    "Notifications that remind staff about mark-release work when applicable.",
])
subheading(doc, "How To Create A Training")
steps(doc, [
    "Open Trainings and click Create New Training.",
    "Enter the training title and description, then choose the affiliated company and training type.",
    "Select the module and device model.",
    "Set the start and end date/time.",
    "Assign one or more trainers.",
    "Select healthcare facilities; use the resulting trainee list to assign participants.",
    "Assign existing device serial numbers or add a custom device when required.",
    "Submit the form, open the new training, and complete its Course, test, attendance, and Settings configuration.",
])
subheading(doc, "How To Bulk Import Trainings")
steps(doc, [
    "Download the training import template from the Trainings page.",
    "Complete the spreadsheet using the expected columns and reference values.",
    "Choose Import Excel and upload the file.",
    "Review validation feedback and correct any unresolved module, title, trainer, healthcare, trainee, date, or device value.",
    "Complete the import and open the created training cards for final configuration.",
])
subheading(doc, "Training Lifecycle Controls")
bullets(doc, [
    "Draft: the training is being prepared.",
    "Active: enrolled trainees can follow the enabled learning flow.",
    "Locked: editing or learner interaction is restricted until the training is unlocked.",
    "Completed: the operational delivery is finished and results/certification records should be verified.",
])

main_heading(doc, "3. Training Workspace Tabs")
label_paragraph(doc, "Purpose: ", "Opening a training displays a tabbed workspace. The visible tabs depend on role, training type, enrollment, lock state, and results-release status.")
add_table(doc, ["Training tab", "Purpose", "Visible to"], [
    ("Stream", "Training overview, trainer identity, sections or announcements/context.", "All roles with access."),
    ("Materials", "Trainee-focused list of assigned learning materials.", "Trainee."),
    ("Tests", "Start available pre, post, certificate, or refresher assessments.", "Trainee."),
    ("Course", "Sections, learning materials, media, tests, and learning sequence.", "All; editing for Admin/Trainer."),
    ("People", "Assigned trainers and enrolled trainees.", "All roles with access."),
    ("My Results", "Released marks, trainer comments, download and certificate actions.", "Trainee after release."),
    ("Marks", "Test responses, scores, practical evaluation, grade/release controls.", "Admin/Trainer."),
    ("Attendance", "Sessions and per-trainee attendance marking/update.", "Admin/Trainer."),
    ("Package", "Generate and download the formal training document package.", "Admin/Trainer."),
    ("Practical Learning Outcomes", "Configure training-specific practical assessment criteria.", "Admin, main training only."),
    ("Settings", "Training metadata, assignments, devices, status and lifecycle controls.", "Admin/Trainer."),
], [1.55, 3.45, 1.5])

subheading(doc, "3.1 Stream")
paragraph(doc, "Use Stream as the training landing page. Review the training identity, trainer or trainers, description, and the current course structure before opening a task-specific tab.")
subheading(doc, "3.2 Materials, Trainee")
steps(doc, [
    "Open Materials from the training workspace.",
    "Choose a section and open the assigned document, image, link, video, or other material.",
    "Complete materials in the order set by the training team when sequencing is important.",
    "Return to Course or Tests to continue the required workflow.",
])
subheading(doc, "3.3 Tests, Trainee")
steps(doc, [
    "Open Tests and review which assessment types are currently available.",
    "Select Start for the required pre-test, post-test, certificate test, or refresher test.",
    "Answer each question and submit the attempt.",
    "Review the immediate result page when the assessment permits it; final training results remain controlled by staff release.",
])
subheading(doc, "3.4 Course")
bullets(doc, [
    "Trainees use Course to follow sections, materials, tests, media, and attendance links.",
    "Admin/Trainer users create or edit sections, add or remove materials, upload media, and organize the learning structure while the training is editable.",
    "Materials may be documents, images, links, videos, or uploaded files. Large media uses chunked upload and should be allowed to finish before leaving the page.",
    "Staff may import sections from another training when the source structure is suitable.",
])
subheading(doc, "How To Add Course Content")
steps(doc, [
    "Open the training and select Course.",
    "Create a section and give it a clear learning-sequence title.",
    "Add material to the section, choose the material type, and provide the title, file, link, or content required by that type.",
    "Save and verify that the item appears in the intended section.",
    "Edit or remove content only while the training state allows changes.",
])
subheading(doc, "3.5 People")
paragraph(doc, "People lists the assigned trainer or trainers and all enrolled trainees. Use it to verify that the expected participants are attached before attendance, tests, grading, or package generation begins.")
subheading(doc, "3.6 My Results, Trainee")
paragraph(doc, "My Results appears only when the trainee has an enrollment and staff has enabled result download. It can show assessment scores, practical outcome results, trainer comments, final status, result download, and certificate availability.")
subheading(doc, "3.7 Marks, Admin/Trainer")
steps(doc, [
    "Open Marks and locate the trainee.",
    "Review submitted test responses and answers when verification is required.",
    "Enter or confirm practical learning outcome scores and trainer comments.",
    "Calculate or confirm the combined result according to the training rules.",
    "Use the appropriate release control only after attendance, tests, and practical outcomes are complete.",
    "For certificate or refresher exceptions, use the documented release/override control and provide the required justification.",
])
subheading(doc, "3.8 Attendance, Admin/Trainer")
steps(doc, [
    "Open Attendance and choose or create the relevant session date.",
    "Mark each trainee Present, Absent, Late, or the status supported by the session.",
    "Use bulk marking when the same status applies to multiple trainees.",
    "Save and re-open the session details to verify the records.",
    "Update attendance in bulk when a completed session needs correction.",
])
subheading(doc, "3.9 Package, Admin/Trainer")
paragraph(doc, "Package prepares formal training records through a background job. Depending on the selected options, the package can include generated PDF forms, letters, attendance or score records, spreadsheets, and a ZIP download.")
steps(doc, [
    "Verify training settings, participants, attendance, marks, device model, and healthcare information.",
    "Open Package and complete the required package fields/options.",
    "Start generation and keep the page open while the job status is checked.",
    "When complete, download the ZIP or requested letter PDF.",
    "Open the generated files and confirm names, dates, participants, scores, and signatures before external use.",
])
subheading(doc, "3.10 Practical Learning Outcomes, Admin/Main Training")
paragraph(doc, "This tab manages the practical criteria attached to a main training. Each aspect should have a clear title, description, and maximum score so the Marks tab can evaluate trainees consistently.")
subheading(doc, "3.11 Training Settings, Admin/Trainer")
bullets(doc, [
    "General training details: title, description, type, module, device model, company, and schedule.",
    "Healthcare, trainer, trainee, and device assignments.",
    "Status, lock/unlock, score-release, enrollment, and other lifecycle controls shown for the role and training state.",
])

main_heading(doc, "4. Trainees")
label_paragraph(doc, "Purpose: ", "Trainees is the staff directory for learner identity, healthcare affiliation, specialization, status, training history, and bulk administration.")
subheading(doc, "Information Displayed")
bullets(doc, [
    "Trainee ID, name, email, IC/passport, telephone, healthcare, designation, specialization, and status.",
    "Auto-managed training values such as serial number, completed-training count, first/latest training, and recertification date where available.",
    "Search, filters, pagination, individual edit, spreadsheet import, and bulk actions.",
])
subheading(doc, "How To Create A Trainee")
steps(doc, [
    "Open Trainees and select Create New Trainee.",
    "Complete Personal information: first name, last name, email, IC/passport, password, and handphone number.",
    "Complete Professional information: healthcare, designation, areas of specialization, and trainee status.",
    "Review the Auto-managed tab; do not attempt to type into fields maintained by training activity.",
    "Submit and verify that the trainee appears in the list.",
])
subheading(doc, "How To Bulk Import Trainees")
steps(doc, [
    "Download the trainee import template.",
    "Enter one trainee per row using valid healthcare and reference values.",
    "Upload the spreadsheet from the Trainees page.",
    "Review row-level validation errors, correct the source file, and upload again if needed.",
    "Confirm imported trainees before assigning them to a training.",
])
subheading(doc, "Bulk Status And Delete Controls")
paragraph(doc, "Select only the intended rows before using bulk status or bulk delete. Status changes can affect whether a trainee may sign in. Deletion removes a learner record and can affect related enrollment data, so verify dependencies and use the least destructive correction available.")

main_heading(doc, "5. Question Bank")
label_paragraph(doc, "Purpose: ", "Question Bank stores reusable multiple-choice questions mapped to test type, objective, and module for assessment delivery.")
subheading(doc, "Information Displayed")
bullets(doc, [
    "Question text, test type, objective, module, answer options, correct answer, and edit/delete actions.",
    "Search and filters for finding questions by their classification.",
    "Excel template download and bulk upload controls.",
])
subheading(doc, "How To Create A Question")
steps(doc, [
    "Open Question Bank and click Create New Question.",
    "Enter the question text.",
    "Choose the test type, objective, and module.",
    "Enter at least options A and B; add C and D when required.",
    "Select the correct answer and save.",
    "Re-open the question and verify the wording and answer mapping before using it in an assessment.",
])
subheading(doc, "How To Bulk Upload Questions")
steps(doc, [
    "Download the question template and review the column instructions.",
    "Enter questions using consistent test-type, objective, module, option, and correct-answer values.",
    "Open Bulk Upload, choose the spreadsheet, and review the preview/validation result.",
    "Correct all reported row errors before completing the upload.",
])

main_heading(doc, "6. Tests And Assessment Attempts")
label_paragraph(doc, "Purpose: ", "Tests deliver pre, post, certificate, and refresher assessments to enrolled trainees and retain attempts for staff review and grading.")
subheading(doc, "Assessment Flow")
bullets(doc, [
    "The training and question configuration determines which assessment types are available.",
    "The trainee starts a test from the Tests tab, answers the presented questions, and submits the attempt.",
    "The system stores the attempt and calculated score.",
    "Admin/Trainer users can inspect responses from Marks before releasing final results.",
])
subheading(doc, "Controls")
bullets(doc, [
    "Do not release a final result solely because a single test is complete; confirm attendance and practical requirements.",
    "Question edits can affect later test generation, so verify the question bank before activating a training.",
    "Certificate and refresher pathways may have separate release conditions and exception controls.",
])

main_heading(doc, "7. Attendance")
label_paragraph(doc, "Purpose: ", "Attendance records presence by training session and enrollment, supports bulk marking, and gives trainees access to their own attendance view.")
subheading(doc, "Information Displayed")
bullets(doc, [
    "Training sessions, session dates, enrolled trainees, public trainee IDs, and attendance status.",
    "Bulk marking and update controls for staff.",
    "A trainee-specific attendance view linked from Course or Tests.",
])
subheading(doc, "How To Correct Attendance")
steps(doc, [
    "Open Attendance and locate the session.",
    "Open session details and compare the attendee list with the source attendance record.",
    "Select the corrected status for each affected trainee.",
    "Save the bulk update and verify the trainee attendance view.",
])

main_heading(doc, "8. Marks, Results And Certificates")
label_paragraph(doc, "Purpose: ", "These functions combine test attempts, practical outcomes, attendance-related readiness, staff comments, grade calculation, result release, downloads, and certificate issuance.")
subheading(doc, "Staff Review Sequence")
steps(doc, [
    "Confirm that all expected test attempts are submitted.",
    "Review test responses when a score requires verification.",
    "Enter practical learning outcome scores and comments.",
    "Calculate or confirm the grade.",
    "Resolve missing attendance or required evidence.",
    "Release grades/scores only when the record is complete.",
    "Enable result download and verify certificate eligibility or issue status.",
])
subheading(doc, "Trainee Result Sequence")
steps(doc, [
    "Open the training after staff release.",
    "Open My Results.",
    "Review test, practical, final result, and trainer comments.",
    "Download the result document when enabled.",
    "Download the certificate when it has been issued and made available.",
])
subheading(doc, "Release Controls")
paragraph(doc, "Some release actions differ by training type and role. Main-training release may require Admin authority, while refresher or certificate exceptions may expose a documented override. Always record the reason and verify the learner record after an override.")

main_heading(doc, "9. Package")
label_paragraph(doc, "Purpose: ", "Package compiles operational evidence into formal files for internal records, customers, healthcare facilities, or audit use.")
subheading(doc, "Package Inputs")
bullets(doc, [
    "Training identity, dates, module, company, device model, healthcare facility, and participants.",
    "Attendance sessions and trainee attendance.",
    "Marks, practical results, and certificate-related information.",
    "Letter/form options and the requested output bundle.",
])
subheading(doc, "Quality Check Before Download")
bullets(doc, [
    "Participant names and IDs match the training enrollment.",
    "Training dates, healthcare, model, device, and module are correct.",
    "Attendance and score data are complete.",
    "Generated PDFs and spreadsheets open correctly before the ZIP is shared.",
])

main_heading(doc, "10. Settings")
label_paragraph(doc, "Purpose: ", "Settings maintains the reference lists used in trainee, question, training, assessment, device, and certificate workflows.")
add_table(doc, ["Settings tab", "Purpose", "Important fields/actions"], [
    ("Modules", "Learning/device module catalog.", "Create, edit, delete; used by trainings and questions."),
    ("Training Titles", "Standardized training-name catalog.", "Create, edit, delete; select during training setup."),
    ("Objectives", "Assessment/learning objective catalog.", "Create, edit, delete; map questions to objectives."),
    ("Healthcare", "Healthcare facility/client records.", "Name, address, CRM linkage, reminder interval, create/edit/delete."),
    ("Areas of Specialization", "Professional specialty options for trainees/users.", "Create, edit, delete."),
    ("Device Serial Numbers", "Device inventory used in training assignment.", "Serial number, model, notes, create/edit/delete."),
    ("Device Models", "Device-model reference catalog.", "Create, edit, delete; used by devices and trainings."),
    ("Practical Learning Outcomes", "Reusable practical assessment aspects.", "Name, description, maximum score, create/edit/delete."),
    ("Users", "Staff account and role administration.", "Admin only; create/edit/delete, role, profile fields, certificate upload."),
], [1.7, 2.25, 2.55])
subheading(doc, "How To Maintain A Reference List")
steps(doc, [
    "Open Settings and select the required reference tab.",
    "Search the list first to avoid creating a duplicate value.",
    "Choose New/Add, complete the required fields, and save.",
    "Edit an existing record when the identity is the same and only its details changed.",
    "Delete only after checking whether the record is already referenced by trainees, questions, trainings, devices, or results.",
])
subheading(doc, "How To Manage Staff Users, Admin")
steps(doc, [
    "Open Settings > Users and click Add User.",
    "Enter identity, email, password, role, position, telephone, specialization, and certificate file as required.",
    "Save and verify that the user can sign in with the intended role.",
    "Use Edit to update details, role, password, status-related data, or certificate information.",
    "Delete only after confirming ownership and historical records that refer to the user.",
])

main_heading(doc, "11. Profile, Registration And Sign-In")
label_paragraph(doc, "Purpose: ", "Profile stores the signed-in user's own information and password. Registration creates a trainee account and links it to healthcare and professional information.")
subheading(doc, "How A Trainee Registers")
steps(doc, [
    "Open Register from the public authentication page.",
    "Complete identity, email, password, IC/passport, healthcare, designation, specialization, device and contact fields shown by the form.",
    "Submit the registration and wait for the resulting trainee status/access rules.",
    "Sign in only while the account status is Active or Registered; contact the administrator if access is blocked.",
])
subheading(doc, "How To Update Profile")
steps(doc, [
    "Open Profile from the sidebar or mobile navigation.",
    "Update the editable personal/professional fields.",
    "Upload a profile picture if required.",
    "Admin/Trainer users may upload or replace their certificate file where the profile exposes that control.",
    "Save and verify the updated profile header and details.",
])
subheading(doc, "How To Change Password")
steps(doc, [
    "Open Profile and find Change Password.",
    "Enter the current password.",
    "Enter and confirm the new password.",
    "Submit, then use the new password at the next sign-in.",
])

main_heading(doc, "12. System Controls And Good Practice")
subheading(doc, "Required Validation")
bullets(doc, [
    "Training dates must form a valid period and required module/title/type assignments must be present.",
    "Trainee email and identity values must be unique where enforced; passwords must meet the minimum length.",
    "Questions require valid classifications, enough answer options, and one correct answer.",
    "Only enrolled trainees should receive training attempts, attendance, grades, results, and certificates.",
    "Result release should follow completed test, practical, and attendance checks.",
    "Uploaded files must use supported formats and should be opened after upload or generation.",
])
subheading(doc, "Recommended Operating Routine")
steps(doc, [
    "Review Dashboard for training status, registrations, recertification, and pending work.",
    "Keep Settings reference data clean before creating trainings, trainees, devices, or questions.",
    "Create or import the training, then verify dates, trainers, healthcare, trainees, and devices.",
    "Prepare sections, materials, tests, and practical outcomes before activating the training.",
    "Verify People before delivery and record attendance during each session.",
    "Review test responses and practical outcomes in Marks.",
    "Release results only after the record is complete, then verify trainee downloads and certificate state.",
    "Generate the Package last, open every output, and retain the verified bundle.",
])
subheading(doc, "Complete Information Flow")
label_paragraph(doc, "LMS flow: ", "CRM/reference settings -> Staff and trainee identities -> Training setup and enrollment -> Course sections/materials/tests -> Attendance and attempts -> Practical evaluation and marks -> Result release -> Certificate and package generation -> Dashboard and profile history.")

footer = section.footer
footer_p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
for run in list(footer_p.runs):
    run._element.getparent().remove(run._element)
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_after = Pt(0)
run = footer_p.add_run("Q-Learning Management System User Guide")
set_font(run, 9)

doc.core_properties.title = "Q-Learning Management System Per Tab Explanation"
doc.core_properties.subject = "Tab explanation and step-by-step user guide"
doc.core_properties.author = "Quick Stop Solution"
doc.core_properties.keywords = "LMS, training, admin, trainer, trainee, user guide"
doc.save(OUTPUT)

if hashlib.sha256(REFERENCE.read_bytes()).hexdigest() != EXPECTED_SHA:
    raise RuntimeError("Reference DOCX was modified during generation")

print(OUTPUT)
